"""Append minimal final-model addenda to the historical course workbooks.

The original course answers stay intact.  The addenda explicitly supersede only
conflicting project-status wording and point readers to the reproducible reports.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
WORKBOOKS = {
    ROOT / "docs/Capstone_Kickoff_Workbook_Serhii_Diadechko.docx": (
        "Final modelling update",
        "This addendum supersedes earlier conflicting project-status wording while preserving the completed course answers and tables as historical context. It does not claim that those historical tables were converted.",
        "The completed evaluation fits the frozen nine-feature hurdle and historical recurrence baseline on T=2010-2019, validates them on T=2020-2021, and uses one frozen final temporal test at T=2022-2024. The nine-feature hurdle has lower held-out all-row MAE and stronger burned-share-mass capture than the baseline, but it underpredicts the unusually high observed mean burned share in T=2024.",
        "The retained output is therefore a continuous comparative research model, not a probability, safety rating, property-level forecast, or purchase recommendation. The buyer-facing capstone output remains historical screening: 1 km mainland grid cells with fire recurrence measured in a 2 km context. The continuous target is burned_share_next_year; burned_next_year remains deferred. See reports/validation/model_final_decision.md.",
    ),
    ROOT / "docs/Repository Setup & Documentation Lab Serhii Diadechko.docx": (
        "Final modelling and repository update",
        "This addendum supersedes earlier conflicting project-status wording while preserving the completed documentation-lab answers and tables as historical context. It does not claim that those historical tables were converted.",
        "The reusable final artifact is a fixed nine-feature hurdle model trained after evaluation on development years T=2010-2021 only. Its seven canonical panel predictors are supplemented by two T-only ERA5-Land monthly-extreme features: maximum monthly mean 2 m temperature and minimum monthly mean layer-1 soil water. The held-out T=2022-2024 years were excluded from that refit.",
        "Final evaluation supports comparative continuous model research but not a buyer-facing recommendation, because the model underpredicts the high-burned T=2024 outcome. The repository keeps Parquet as the canonical analytical table, GeoPackages for validated historical/GIS outputs, and the QGIS project as a presentation of historical evidence rather than a model recommendation. See reports/validation/final_temporal_test_2022_2024.md and reports/validation/model_final_decision.md.",
    ),
}


def append_addendum(path: Path, values: tuple[str, str, str, str]) -> bool:
    document = Document(path)
    heading, *paragraphs = values
    if any(item.text.strip() == heading for item in document.paragraphs):
        return False
    document.add_page_break()
    document.add_paragraph(heading, style="Heading 1")
    for text in paragraphs:
        document.add_paragraph(text)
    document.save(path)
    return True


if __name__ == "__main__":
    for file_path, content in WORKBOOKS.items():
        print(f"{file_path.name}: {'updated' if append_addendum(file_path, content) else 'already current'}")
