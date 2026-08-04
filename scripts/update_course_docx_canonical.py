"""Focused canonical-design addendum for the two retained course DOCX artifacts."""
from pathlib import Path
from docx import Document
from docx.shared import Pt
ROOT=Path(__file__).resolve().parents[1]
COMMON=("This addendum supersedes earlier conflicting design and source wording while preserving the original completed course answers as historical context. It does not claim that the historical tables themselves were converted. One EPSG:3763 1 km cell-year is the only analytical unit; 2 km is an outward context buffer. Train T=2015–2019, validate T=2020–2021, and reserve T=2022–2024 (outcomes 2023–2025) for final temporal test. The 2023→2024 artifact is a data-contract/pipeline feasibility pilot only. ")
WORKBOOK=(COMMON+"The canonical schema stores identifiers and geometry separately, with seven predictors: built_up_share, forest_shrub_share_2km, mean_slope_2km, fire_years_previous_10y_2km, warm_season_mean_2m_temperature_c, warm_season_total_precipitation_mm, and warm_season_mean_soil_water_layer1. The continuous target is burned_share_next_year; burned_next_year and predicted wildfire probability are deferred until target-distribution review and a documented threshold. Copernicus CLC governs MVP land cover and Copernicus DEM GLO-30 governs static 2 km slope context.")
LAB=(COMMON+"The canonical analytical schema is seven predictors plus the continuous burned_share_next_year target, with identifiers and geometry separate. The predictors include layer-1 soil water alongside JJAS temperature and day-weighted precipitation. Copernicus CLC is the MVP land-cover source and Copernicus DEM GLO-30 is the terrain source for mean_slope_2km. The full temporal split is training T=2015–2019, validation T=2020–2021, and final test T=2022–2024; 2023→2024 is feasibility-pilot work only.")
for name in ['Capstone_Kickoff_Workbook_Serhii_Diadechko.docx','Repository Setup & Documentation Lab Serhii Diadechko.docx']:
 d=Document(ROOT/'docs'/name); text=WORKBOOK if name.startswith('Capstone') else LAB
 found=False
 for i,p in enumerate(d.paragraphs):
  if p.text.strip()=='Canonical design update' and i+1<len(d.paragraphs):
   d.paragraphs[i+1].text=text; found=True; break
 if not found:
  d.add_heading('Canonical design update',level=1); d.add_paragraph(text)
 d.save(ROOT/'docs'/name)
